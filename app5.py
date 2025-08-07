import streamlit as st
from docx import Document
import openai
import fitz  # PyMuPDF
from io import BytesIO
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

# Set up the app
st.title("Academic Application Generator")
st.subheader("Create Motivation & Cover Letters (Download in Word or PDF)")

# Session state
if "generated_letter" not in st.session_state:
    st.session_state.generated_letter = ""
if "cover_letter" not in st.session_state:
    st.session_state.cover_letter = ""

# Sidebar
with st.sidebar:
    st.header("Configuration")
    openai_api_key = st.text_input("Enter OpenAI API key:", type="password")
    st.markdown("[Get OpenAI API key](https://platform.openai.com/account/api-keys)")

def main():
    with st.expander("Step 1: Program/Job Details", expanded=True):
        job_ad = st.text_area("Paste advertisement:", height=200)

    with st.expander("Step 2: Committee/HR Profiles"):
        recruiter_profiles = st.text_area("Paste decision-makers' profiles:", height=150)

    with st.expander("Step 3: Upload CV"):
        cv_file = st.file_uploader("Upload CV (Word/Text/PDF):", type=["docx", "txt", "pdf"])
        cv_text = ""
        if cv_file:
            if cv_file.name.endswith('.docx'):
                doc = Document(cv_file)
                cv_text = "\n".join([para.text for para in doc.paragraphs])
            elif cv_file.name.endswith('.pdf'):
                pdf = fitz.open(stream=cv_file.read(), filetype="pdf")
                cv_text = "\n".join([page.get_text() for page in pdf])
            else:
                cv_text = cv_file.getvalue().decode()

    with st.expander("Step 4: Interests & Goals"):
        interests = st.text_area("Your specific interests and goals:", height=150)

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Generate Motivation Letter"):
            generate_letter(job_ad, recruiter_profiles, cv_text, interests, "motivation")

    with col2:
        if st.button("Generate Cover Letter"):
            generate_letter(job_ad, recruiter_profiles, cv_text, interests, "cover")

    if st.session_state.generated_letter:
        show_letter("Motivation Letter", st.session_state.generated_letter, "motivation_letter")

        if st.button("Convert to Cover Letter"):
            convert_letter_type()

    if st.session_state.cover_letter:
        show_letter("Cover Letter", st.session_state.cover_letter, "cover_letter")

def generate_letter(job_ad, profiles, cv, interests, letter_type):
    if not openai_api_key:
        st.error("API key required")
        return
    if not job_ad or not cv:
        st.error("Missing required fields")
        return

    prompt = create_prompt(job_ad, profiles, cv, interests, letter_type)

    try:
        openai.api_key = openai_api_key
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": get_system_message(letter_type)},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=1200 if letter_type == "motivation" else 800
        )

        content = response.choices[0].message.content
        if letter_type == "motivation":
            st.session_state.generated_letter = content
        else:
            st.session_state.cover_letter = content

    except Exception as e:
        st.error(f"Generation error: {str(e)}")

def convert_letter_type():
    if not st.session_state.generated_letter:
        st.error("Generate motivation letter first")
        return

    try:
        openai.api_key = openai_api_key
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "Professional document converter"},
                {"role": "user", "content": f"Convert this academic motivation letter into a professional cover letter:\n\n{st.session_state.generated_letter}"}
            ],
            temperature=0.6,
            max_tokens=700
        )
        st.session_state.cover_letter = response.choices[0].message.content
    except Exception as e:
        st.error(f"Conversion error: {str(e)}")

def create_prompt(job_ad, profiles, cv, interests, letter_type):
    return f"""
    Act as a {'graduate admissions committee member' if letter_type == 'motivation' else 'professional career advisor'}
    to create a {'motivation' if letter_type == 'motivation' else 'cover'} letter using:

    1. Advertisement: {job_ad}
    2. Decision-makers' Profiles: {profiles}
    3. Applicant CV: {cv}
    4. Applicant Interests: {interests}

    Requirements:
    - Match key skills with {'program' if letter_type == 'motivation' else 'job'} requirements
    - Align goals with {'research objectives' if letter_type == 'motivation' else 'position goals'}
    - Demonstrate strong interest and fit
    - {"Academic tone with research focus" if letter_type == 'motivation' else "Professional business format"}
    - Length: {"500-700 words" if letter_type == 'motivation' else "300-500 words"}
    """

def get_system_message(letter_type):
    return "Expert academic writer" if letter_type == "motivation" else "Professional career consultant"

def show_letter(title, content, filename_base):
    st.subheader(title)
    st.write(content)

    # Plain text download
    st.download_button(
        label=f"Download {title} (.txt)",
        data=content,
        file_name=f"{filename_base}.txt",
        mime="text/plain"
    )

    # Word (.docx) download
    word_file = create_docx(content)
    st.download_button(
        label=f"Download {title} (.docx)",
        data=word_file,
        file_name=f"{filename_base}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # PDF download
    pdf_file = create_pdf(content)
    st.download_button(
        label=f"Download {title} (.pdf)",
        data=pdf_file,
        file_name=f"{filename_base}.pdf",
        mime="application/pdf"
    )

def create_docx(text):
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_pdf(text):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer)
    styles = getSampleStyleSheet()
    flow = [Paragraph(line, styles["Normal"]) for line in text.split("\n") if line.strip()]
    doc.build(flow)
    buffer.seek(0)
    return buffer

if __name__ == "__main__":
    main()
